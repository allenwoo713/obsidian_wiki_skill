"""extract_assets.py 测试：调度与落盘。"""
import io
import sys
from pathlib import Path
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from extract_assets import extract, UnsupportedFormat
from models import ParsedDoc


def _valid_png_bytes():
    """用 Pillow 生成有效 PNG 用于测试。"""
    from PIL import Image
    img = Image.new("RGB", (4, 4), color=(0, 100, 200))
    img.putpixel((1, 1), (255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


_DEFAULT_PNG = _valid_png_bytes()


def _make_test_docx(tmp_path, image_bytes=None):
    from docx import Document
    from docx.shared import Inches
    if image_bytes is None:
        image_bytes = _DEFAULT_PNG
    doc = Document()
    doc.add_paragraph("正文")
    img = tmp_path / "t.png"
    img.write_bytes(image_bytes)
    doc.add_picture(str(img), width=Inches(1))
    doc.add_paragraph("图1 示意图")
    docx_path = tmp_path / "t.docx"
    doc.save(str(docx_path))
    return docx_path, img.read_bytes()


def test_extract_returns_parsed_doc(tmp_path):
    docx_path, _ = _make_test_docx(tmp_path)
    assets_dir = tmp_path / "assets"
    parsed = extract(docx_path, assets_dir)
    assert isinstance(parsed, ParsedDoc)
    assert len(parsed.images) == 1
    assert "正文" in parsed.text


def test_extract_writes_image_to_assets(tmp_path):
    docx_path, img_bytes = _make_test_docx(tmp_path)
    assets_dir = tmp_path / "assets"
    parsed = extract(docx_path, assets_dir)
    assert assets_dir.exists()
    written = list(assets_dir.iterdir())
    assert len(written) == 1
    assert written[0].read_bytes() == img_bytes


def test_extract_dedup_same_sha256(tmp_path):
    """同一 docx 两张相同图片只落盘一份。"""
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    img_bytes = _valid_png_bytes()
    img = tmp_path / "same.png"
    img.write_bytes(img_bytes)
    doc.add_picture(str(img), width=Inches(1))
    doc.add_paragraph("图1")
    doc.add_picture(str(img), width=Inches(1))
    doc.add_paragraph("图2")
    docx_path = tmp_path / "dup.docx"
    doc.save(str(docx_path))
    assets_dir = tmp_path / "assets"
    parsed = extract(docx_path, assets_dir)
    assert len(parsed.images) == 2
    written = list(assets_dir.iterdir())
    assert len(written) == 1


def test_extract_unsupported_format(tmp_path):
    f = tmp_path / "x.xyz"
    f.write_text("?", encoding="utf-8")
    with pytest.raises(UnsupportedFormat):
        extract(f, tmp_path / "assets")


def _make_minimal_pdf(tmp_path):
    """构造最小 PDF 用于 extract 测试。"""
    import fitz
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), "extract test body", fontsize=12)
    pdf_path = tmp_path / "e.pdf"
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


def test_extract_pdf_default_local_backend(tmp_path, monkeypatch):
    """未设环境变量 → 走本地 PdfParser。"""
    monkeypatch.delenv("PDF_PARSER_BACKEND", raising=False)
    monkeypatch.delenv("FIRECRAWL_API_KEY", raising=False)
    pdf_path = _make_minimal_pdf(tmp_path)
    parsed = extract(pdf_path, tmp_path / "assets")
    assert "extract test body" in parsed.text


def test_extract_pdf_firecrawl_backend_no_key_falls_back(tmp_path, monkeypatch):
    """设了 BACKEND=firecrawl 但无 API key → 回退本地。"""
    monkeypatch.setenv("PDF_PARSER_BACKEND", "firecrawl")
    monkeypatch.delenv("FIRECRAWL_API_KEY", raising=False)
    pdf_path = _make_minimal_pdf(tmp_path)
    parsed = extract(pdf_path, tmp_path / "assets")
    assert "extract test body" in parsed.text


def test_extract_pdf_firecrawl_import_error_falls_back(tmp_path, monkeypatch):
    """设了 BACKEND + KEY，但 firecrawl_pdf_parser 导入失败 → 调度层回退。"""
    monkeypatch.setenv("PDF_PARSER_BACKEND", "firecrawl")
    monkeypatch.setenv("FIRECRAWL_API_KEY", "fake")
    import sys
    monkeypatch.setitem(sys.modules, "parsers.firecrawl_pdf_parser", None)
    pdf_path = _make_minimal_pdf(tmp_path)
    parsed = extract(pdf_path, tmp_path / "assets")
    assert "extract test body" in parsed.text
