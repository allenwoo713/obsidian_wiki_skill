"""DocxParser 测试：用 python-docx 构造测试 docx。"""
import sys
import hashlib
import io
from pathlib import Path
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from parsers.docx_parser import DocxParser
from parsers.base import ParseResult


def _valid_png_bytes(r=0, g=0, b=0):
    """用 Pillow 生成有效 PNG（基准彩图，双色）用于 md5/SHA 区分。"""
    from PIL import Image
    img = Image.new("RGB", (4, 4), color=(r, g, b))
    # 加对比像素确保编码内容不同
    img.putpixel((1, 1), (255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


_DEFAULT_PNG = _valid_png_bytes(0, 100, 200)


def _make_docx_with_image(tmp_path, image_bytes=None, caption="图1 测试示意图"):
    """用 python-docx 构造含一张图片+图注的测试 docx。"""
    from docx import Document
    from docx.shared import Inches
    if image_bytes is None:
        image_bytes = _DEFAULT_PNG
    doc = Document()
    doc.add_paragraph("文档开头正文。")
    img_path = tmp_path / "test_img.png"
    img_path.write_bytes(image_bytes)
    doc.add_picture(str(img_path), width=Inches(1))
    if caption:
        doc.add_paragraph(caption)
    doc.add_paragraph("图片后续正文。")
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))
    return docx_path, image_bytes


def test_docx_parser_returns_parse_result(tmp_path):
    docx_path, _ = _make_docx_with_image(tmp_path)
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert isinstance(result, ParseResult)
    assert "文档开头正文" in result.text
    assert "图片后续正文" in result.text


def test_docx_parser_extracts_image(tmp_path):
    img_bytes = _valid_png_bytes(100, 0, 0)
    docx_path, _ = _make_docx_with_image(tmp_path, image_bytes=img_bytes)
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert len(result.images) == 1
    ref = result.images[0]
    # 文件名含 _img
    assert "_img" in ref.filename or "img01" in ref.filename
    assert ref.rel_path.startswith("assets/")
    assert len(ref.sha256) == 64
    assert ref.source_media_name.startswith("image")
    assert ref.page_or_section == "body"
    # _image_bytes 与 images 同序
    assert len(result._image_bytes) == 1
    assert hashlib.sha256(result._image_bytes[0]).hexdigest() == ref.sha256


def test_docx_parser_image_renders_as_obsidian_embed(tmp_path):
    """图片以 ![[filename]] Obsidian 嵌入格式出现，图注作为可读文本，非 {{IMG|}} 占位符。"""
    docx_path, _ = _make_docx_with_image(tmp_path, caption="图1 测试示意图")
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert "{{IMG|" not in result.text
    assert "![" in result.text
    assert "图1 测试示意图" in result.text


def test_docx_parser_caption_extracted(tmp_path):
    docx_path, _ = _make_docx_with_image(tmp_path, caption="图1 方位角 FOV ±45° 示意图")
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert len(result.images) == 1
    assert "图1" in result.images[0].caption
    assert "±45°" in result.images[0].caption


def test_docx_parser_no_caption(tmp_path):
    docx_path, _ = _make_docx_with_image(tmp_path, caption="")
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert len(result.images) == 1
    assert result.images[0].caption == ""


def test_docx_parser_multiple_images(tmp_path):
    """构造含两张图片的 docx。"""
    from docx import Document
    from docx.shared import Inches
    img1_bytes = _valid_png_bytes(50, 0, 0)
    img2_bytes = _valid_png_bytes(0, 50, 0)
    doc = Document()
    doc.add_paragraph("正文1")
    img1 = tmp_path / "i1.png"
    img1.write_bytes(img1_bytes)
    doc.add_picture(str(img1), width=Inches(1))
    doc.add_paragraph("图1 第一张")
    img2 = tmp_path / "i2.png"
    img2.write_bytes(img2_bytes)
    doc.add_picture(str(img2), width=Inches(1))
    doc.add_paragraph("图2 第二张")
    docx_path = tmp_path / "multi.docx"
    doc.save(str(docx_path))
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert len(result.images) == 2
    assert result.images[0].sha256 != result.images[1].sha256


def test_docx_parser_table_surrounded_by_blank_lines(tmp_path):
    """表格 markdown 前后应有空行，确保 Obsidian 正确渲染。"""
    from docx import Document
    doc = Document()
    doc.add_paragraph("表格前文本")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Value1"
    table.cell(1, 1).text = "Value2"
    doc.add_paragraph("表格后文本")
    docx_path = tmp_path / "table.docx"
    doc.save(str(docx_path))
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert "| Header1" in result.text
    assert "| Value2" in result.text
    # 表格前有空行
    assert "\n\n| Header1" in result.text
    # 表格后有空行
    assert "Value2 |\n\n" in result.text


def test_docx_parser_heading_converts_to_markdown(tmp_path):
    """Word heading 样式应转为 markdown #/## 层级。"""
    from docx import Document
    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("正文内容")
    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("更多正文")
    docx_path = tmp_path / "heading.docx"
    doc.save(str(docx_path))
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert "# 一级标题" in result.text
    assert "## 二级标题" in result.text


def test_docx_parser_no_duplicate_caption(tmp_path):
    """图注不应重复出现（占位符释放的 caption 与原文档 caption 段落去重）。"""
    docx_path, _ = _make_docx_with_image(tmp_path, caption="图1 测试示意图")
    parser = DocxParser()
    result = parser.parse(docx_path)
    assert result.text.count("图1 测试示意图") == 1
