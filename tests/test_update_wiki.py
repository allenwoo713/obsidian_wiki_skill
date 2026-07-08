"""update_wiki.py 测试。"""
import json
from pathlib import Path
import pytest
import sys

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from update_wiki import scan_sources, diff_manifest, SourceState


def _write_src(raw, name, content="doc content"):
    f = raw / name
    f.parent.mkdir(parents=True, exist_ok=True)
    f.write_text(content, encoding="utf-8")


def test_scan_sources(tmp_path):
    raw = tmp_path / "Raw" / "sources"
    _write_src(raw, "a/a.md", "# A\nbody")
    _write_src(raw, "b/b.docx", "fake")
    _write_src(raw, "c/c.pdf", "fake")
    docs = scan_sources(raw)
    suffixes = sorted(d.suffix for d in docs)
    assert ".docx" in suffixes
    assert ".md" in suffixes
    assert ".pdf" in suffixes


def test_diff_new_file(tmp_path):
    raw = tmp_path / "Raw" / "sources"
    _write_src(raw, "a.md", "# A\nbody")
    docs = scan_sources(raw)
    manifest = {"entries": {}}
    diff = diff_manifest(docs, manifest)
    assert len(diff["new"]) == 1
    assert diff["new"][0].path.name == "a.md"
    assert len(diff["modified"]) == 0
    assert len(diff["unchanged"]) == 0


def test_diff_unchanged(tmp_path):
    raw = tmp_path / "Raw" / "sources"
    _write_src(raw, "a.md", "# A\nbody")
    docs = scan_sources(raw)
    from parse_sources import compute_sha256
    sha = compute_sha256(raw / "a.md")
    manifest = {"entries": {str(raw / "a.md"): {"sha256": sha, "status": "processed"}}}
    diff = diff_manifest(docs, manifest)
    assert len(diff["unchanged"]) == 1
    assert len(diff["new"]) == 0


def test_diff_modified(tmp_path):
    raw = tmp_path / "Raw" / "sources"
    _write_src(raw, "a.md", "# A\nnew body")
    docs = scan_sources(raw)
    manifest = {"entries": {str(raw / "a.md"): {"sha256": "old_sha", "status": "processed"}}}
    diff = diff_manifest(docs, manifest)
    assert len(diff["modified"]) == 1


def _create_minimal_png(path):
    """生成一个最小的合法 1x1 红色 PNG。"""
    import struct, zlib
    sig = b'\x89PNG\r\n\x1a\n'
    ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
    ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xffffffff
    ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
    raw_data = b'\x00\xff\x00\x00'
    compressed = zlib.compress(raw_data)
    idat_crc = zlib.crc32(b'IDAT' + compressed) & 0xffffffff
    idat = struct.pack('>I', len(compressed)) + b'IDAT' + compressed + struct.pack('>I', idat_crc)
    iend_crc = zlib.crc32(b'IEND') & 0xffffffff
    iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
    path.write_bytes(sig + ihdr + idat + iend)


def test_extract_images_for_diff_new(tmp_path):
    from docx import Document; from docx.shared import Inches
    raw = tmp_path / "Raw"; raw.mkdir()
    doc = Document(); img = tmp_path / "t.png"
    _create_minimal_png(img)
    doc.add_picture(str(img), width=Inches(1)); doc.add_paragraph("图1")
    dp = raw / "t.docx"; doc.save(str(dp))
    assets = tmp_path / "assets"
    from update_wiki import extract_images_for_diff
    im = extract_images_for_diff([dp], [], assets)
    assert len(im) == 1; assert im[0]["filename"].endswith("_img01.png")
    # assets 目录被创建且含图片
    assert assets.exists()


def test_write_source_fulltext_no_images_section(tmp_path):
    """write_source_fulltext 不生成 AUTO-GENERATED IMAGES 区（FULLTEXT 区已嵌入图片）。"""
    from update_wiki import write_source_fulltext
    from models import ParsedDoc, ImageRef

    proj = tmp_path
    source_path = proj / "Raw" / "sources" / "test.docx"
    source_path.parent.mkdir(parents=True, exist_ok=True)
    source_path.write_bytes(b"fake")

    img_ref = ImageRef(
        filename="test_img01.png", rel_path="assets/test_img01.png",
        caption="图1 测试", source_media_name="image1.png",
        sha256="abc", page_or_section="body",
    )
    parsed = ParsedDoc(
        path=source_path, title="Test Doc",
        text="正文\n![[test_img01.png]]  \n图1 测试",
        tables=[], sha256="def", doc_type="docx", images=[img_ref],
    )
    page = write_source_fulltext(proj, source_path, parsed)
    content = page.read_text(encoding="utf-8")
    # 不应有 AUTO-GENERATED IMAGES 标记
    assert "BEGIN AUTO-GENERATED IMAGES" not in content
    assert "END AUTO-GENERATED IMAGES" not in content
    # 应有 AUTO-GENERATED FULLTEXT 标记
    assert "BEGIN AUTO-GENERATED FULLTEXT" in content


def test_write_source_fulltext_cleans_legacy_images(tmp_path):
    """write_source_fulltext 清理旧的 AUTO IMAGES 标记区 + LLM 管理区'## 文档内嵌图片'段落。"""
    from update_wiki import write_source_fulltext
    from models import ParsedDoc

    proj = tmp_path
    source_path = proj / "Raw" / "sources" / "test.docx"
    source_path.parent.mkdir(parents=True, exist_ok=True)
    source_path.write_bytes(b"fake")

    page_path = proj / "Wiki" / "sources" / "test.md"
    page_path.parent.mkdir(parents=True, exist_ok=True)
    old_content = (
        "---\ntype: source-summary\ntitle: \"Test\"\n"
        'sources: ["Raw/sources/test.docx"]\nupdated: 2026-06-29\n---\n\n'
        "# Test\n\n## 文档信息\n\n## 核心内容摘要\n\n旧摘要。\n\n"
        "## 文档内嵌图片\n\n![[test_img01.png]] ![[test_img02.png]]\n\n"
        "<!-- BEGIN AUTO-GENERATED FULLTEXT -->\n## 全文内容\n\n旧全文。\n<!-- END AUTO-GENERATED FULLTEXT -->\n\n"
        "<!-- BEGIN AUTO-GENERATED IMAGES -->\n## 文档内嵌图片\n\n![[test_img01.png]]\n<!-- END AUTO-GENERATED IMAGES -->\n"
    )
    page_path.write_text(old_content, encoding="utf-8")

    parsed = ParsedDoc(
        path=source_path, title="Test", text="新全文",
        tables=[], sha256="abc", doc_type="docx", images=[],
    )
    page = write_source_fulltext(proj, source_path, parsed)
    content = page.read_text(encoding="utf-8")
    # "## 文档内嵌图片" 完全消失（旧的 LLM 区 + 旧的 AUTO 区都清理）
    assert "## 文档内嵌图片" not in content
    # AUTO-GENERATED IMAGES 标记被移除
    assert "BEGIN AUTO-GENERATED IMAGES" not in content
    assert "END AUTO-GENERATED IMAGES" not in content
    # FULLTEXT 区保留且更新
    assert "BEGIN AUTO-GENERATED FULLTEXT" in content
    assert "新全文" in content
