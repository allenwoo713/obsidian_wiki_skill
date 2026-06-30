"""parse_sources.py 测试。"""
import hashlib
from pathlib import Path
import pytest
import sys

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from models import ParsedDoc
from parse_sources import parse_docx, parse_pdf, parse_markdown, parse_file, compute_sha256


def test_compute_sha256(tmp_path):
    f = tmp_path / "x.txt"
    f.write_text("hello", encoding="utf-8")
    expected = hashlib.sha256(b"hello").hexdigest()
    assert compute_sha256(f) == expected


def test_parse_markdown(tmp_path):
    f = tmp_path / "note.md"
    f.write_text("# Title\n\n正文段落。", encoding="utf-8")
    doc = parse_markdown(f)
    assert doc.doc_type == "md"
    assert doc.title == "Title"
    assert "正文段落" in doc.text
    assert doc.tables == []
    assert len(doc.sha256) == 64


def test_parse_file_dispatch(tmp_path):
    f = tmp_path / "n.md"
    f.write_text("# H\nbody", encoding="utf-8")
    doc = parse_file(f)
    assert doc.title == "H"


def test_parse_file_unsupported(tmp_path):
    f = tmp_path / "x.xyz"
    f.write_text("?", encoding="utf-8")
    with pytest.raises(ValueError, match="unsupported"):
        parse_file(f)


def _make_minimal_png() -> bytes:
    """生成最小的合法 1x1 白色 PNG。"""
    import struct
    import zlib

    def _chunk(ctype, data):
        raw = ctype + data
        return struct.pack(">I", len(data)) + raw + struct.pack(">I", zlib.crc32(raw) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\xff\xff")  # filter none + RGB（白）
    return b"\x89PNG\r\n\x1a\n" + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


def test_parse_docx_with_assets_extracts_images(tmp_path):
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_paragraph("正文")
    img = tmp_path / "t.png"
    img.write_bytes(_make_minimal_png())
    doc.add_picture(str(img), width=Inches(1))
    doc.add_paragraph("图1 示意")
    docx_path = tmp_path / "t.docx"
    doc.save(str(docx_path))
    from parse_sources import parse_docx
    assets_dir = tmp_path / "assets"
    parsed = parse_docx(docx_path, assets_dir=assets_dir)
    assert len(parsed.images) == 1
    assert assets_dir.exists()


def test_parse_docx_no_assets_backward_compat(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("正文")
    docx_path = tmp_path / "t.docx"
    doc.save(str(docx_path))
    from parse_sources import parse_docx
    parsed = parse_docx(docx_path)
    assert parsed.images == []
    assert "正文" in parsed.text
