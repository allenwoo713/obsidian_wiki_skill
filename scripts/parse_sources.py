"""源文档解析：docx / pdf / pptx / 多格式 Office / md / txt → ParsedDoc。"""
from __future__ import annotations
import hashlib
from pathlib import Path
from typing import List, Optional
import re

from models import ParsedDoc


# 需要 extract_assets 进行解析的二进制/复杂格式（含图片提取）
_BINARY_FORMATS = {
    ".docx",
    ".pdf",
    ".pptx",
    ".doc",
    ".ppt",
    ".xls",
    ".xlsx",
    ".html",
    ".htm",
}


def compute_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _extract_title(text: str, fallback: str) -> str:
    m = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    return m.group(1).strip() if m else fallback


def parse_markdown(path: Path) -> ParsedDoc:
    text = path.read_text(encoding="utf-8", errors="replace")
    title = _extract_title(text, path.stem)
    return ParsedDoc(
        path=path, title=title, text=text, tables=[],
        sha256=compute_sha256(path), doc_type="md",
    )


def parse_txt(path: Path) -> ParsedDoc:
    text = path.read_text(encoding="utf-8", errors="replace")
    title = _extract_title(text, path.stem)
    return ParsedDoc(
        path=path, title=title, text=text, tables=[],
        sha256=compute_sha256(path), doc_type="txt",
    )


def parse_docx(path: Path, assets_dir: Path = None) -> ParsedDoc:
    """DOCX 解析：优先走 extract_assets（提取图片），否则纯文本。"""
    if assets_dir is not None:
        from extract_assets import extract
        return extract(path, assets_dir)
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    tables: List[List[List[str]]] = []
    for t in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
        tables.append(rows)
    if tables:
        text += "\n\n[表格]\n"
        for i, t in enumerate(tables):
            text += f"\n表 {i+1}:\n"
            for row in t:
                text += " | ".join(row) + "\n"
    title = _extract_title(text, path.stem)
    return ParsedDoc(
        path=path, title=title, text=text, tables=tables,
        sha256=compute_sha256(path), doc_type="docx",
    )


def parse_pdf(path: Path, assets_dir: Path = None, is_sensitive: Optional[bool] = None) -> ParsedDoc:
    """PDF 解析：委托 extract_assets，支持敏感文档本地解析。"""
    if assets_dir is not None:
        from extract_assets import extract
        return extract(path, assets_dir, is_sensitive=is_sensitive)
    # 无 assets_dir 时的纯文本兜底
    text = ""
    try:
        import fitz
        with fitz.open(str(path)) as doc:
            for page in doc:
                text += (page.get_text() or "") + "\n"
    except Exception:
        text = ""
    title = _extract_title(text, path.stem)
    return ParsedDoc(
        path=path, title=title, text=text, tables=[],
        sha256=compute_sha256(path), doc_type="pdf",
    )


_PARSERS = {
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_txt,
}


def parse_file(path: Path, assets_dir: Path = None, is_sensitive: Optional[bool] = None) -> ParsedDoc:
    """统一入口解析单个源文档。

    Args:
        path: 源文件路径。
        assets_dir: 图片输出目录。二进制/复杂格式必须提供。
        is_sensitive: 仅对 PDF 有效。True 走 MinerU Local，False/None 走 Cloud。
    """
    suffix = path.suffix.lower()

    if suffix in (".md", ".markdown", ".txt"):
        parser = _PARSERS[suffix]
        return parser(path)

    if suffix in _BINARY_FORMATS:
        if assets_dir is None:
            raise ValueError(f"{suffix} 解析需要提供 assets_dir 以提取图片/表格")
        from extract_assets import extract
        return extract(path, assets_dir, is_sensitive=is_sensitive)

    raise ValueError(f"unsupported file type: {suffix} ({path})")
